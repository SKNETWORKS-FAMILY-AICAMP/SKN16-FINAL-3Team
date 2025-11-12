import docx

doc = docx.Document('temp.docx')

# Table 0
table0_data = [
    ('\ud56d\ubaa9\uba85', '\uc124\uba85', '\uc608\uc2dc'),
    ('document_id', '\ubc95\ub839/\uc0c1\ud488 \uace0\uc720 \uc2dd\ubcc4\uc790', 'LAW_2024_082'),
    ('source', '\ucd9c\ucc98 \uae30\uad00', '\uad6d\uac00\ubc95\ub839\uc815\ubcf4\uc13c\ud130'),
    ('title', '\ubc95\ub839 \ub610\ub294 \uc0c1\ud488\uba85', '\uc740\ud589\ubc95 \uc81c17\uc870(\uacb8\uc601\uc5c5\ubb34\uc758 \ubc94\uc704)'),
    ('content_preview', '\uc694\uc57d\ub41c \ubcf8\ubb38', '\uacb8\uc601\uc5c5\ubb34 \ubc94\uc704 \ud575\uc2ec \ubb38\uc7a5 \uc694\uc57d'),
    ('effective_date', '\uc2dc\ud589/\uacf5\uc2dc \uc77c\uc790', '2024-12-31'),
]
for row, values in zip(doc.tables[0].rows, table0_data):
    for cell, value in zip(row.cells, values):
        cell.text = value

# Table 1
table1_data = [
    ('\ub2e8\uacc4', '\uc8fc\uc694 \uc791\uc5c5', '\ube44\uace0', '\uc0ac\uc6a9 \ub77c\uc774\ube0c\ub7ec\ub9ac/\ub3c4\uad6c'),
    ('\uc218\uc9d1', 'API\ub97c \ud1b5\ud574 \ubc95\ub839\ubd80\ud488 \ub370\uc774\ud130 \uc218\uc9d1', '', 'requests, pandas'),
    ('\uc815\uc81c', '\uc911\ubcf5\uac12 \uc81c\uac70 \ubc0f \uc77c\uc790 \ud615\uc2dd \ud1b5\uc77c', '', 'pandas, numpy'),
    ('\uc815\uaddc\ud654', '\ubb38\uc7a5 \ubd84\ud560\uacfc \ud1a0\ud070 \ubd84\ub9ac \uc218\ud589', '', 'KoNLPy, nltk'),
    ('\ud0dc\uadf8', '\ubc95\ub839/\uc0c1\ud488 \uc720\ud615 \ud0dc\uadf8 \ubc0f \uc784\ubca0\ub529 \uc900\ube44', '', 'scikit-learn, sentence-transformers'),
    ('\ubd84\ud560', '\ud559\uc2b5/\uac80\uc99d \uc138\ud2b8 \ubd84\ud560', 'train:valid = 8:2', 'sklearn.model_selection.train_test_split'),
]
for row, values in zip(doc.tables[1].rows, table1_data):
    for cell, value in zip(row.cells, values):
        cell.text = value

# Table 2
table2_data = [
    ('\ud56d\ubaa9', '\uacb0\uc998 \ube44\uc728', '\uc870\uce58'),
    ('effective_date', '3.8%', '\ubbf8\ud655\uc778 \u2192 \uc815\ubcf4\uc5c6\uc74c \ub300\uccb4'),
    ('content_preview', '1.2%', '\uacb0\uc998 \ud589 \uc81c\uac70'),
]
for row, values in zip(doc.tables[2].rows, table2_data):
    for cell, value in zip(row.cells, values):
        cell.text = value

# Table 3
table3_data = [
    ('\uc9c0\ud45c', '\uae30\uc900', '\uc870\uce58', '\uc218\ub7c9'),
    ('content_length', '> 4,000\uc790', '\uc694\uc57d \ud6c4 4,000\uc790 \uc774\ud558\ub85c \uc808\ub2e8', '37\uac74'),
    ('interest_rate', '< 0 \ub610\ub294 > 20', '\uc0c1\ud488 \ub2f4\ub2f9\uc790 \uac80\uc99d \ud50c\ub798\uadf8', '12\uac74'),
]
for row, values in zip(doc.tables[3].rows, table3_data):
    for cell, value in zip(row.cells, values):
        cell.text = value

# Table 4
table4_data = [
    ('\uad6c\ubd84', '\uac74\uc218'),
    ('\ud559\uc2b5 \ub370\uc774\ud130', '1,310\uac74'),
    ('\uac80\uc99d \ub370\uc774\ud130', '327\uac74'),
]
for row, values in zip(doc.tables[4].rows, table4_data):
    for cell, value in zip(row.cells, values):
        cell.text = value

# Table 5
table5_data = [
    ('\ubcc0\uacbd\uc77c', '\uc791\uc131\uc790', '\ubcc0\uacbd \ub0b4\uc6a9', '\ube44\uace0'),
    ('2025-11-10', '\uae40\ud558\ub298', '\ub370\uc774\ud130 \uc815\uc81c \ub85c\uc9c1 \ucd5c\uc2e0 \ubc95\ub839 \ubc18\uc751', 'tag \uae30\uc900 \uc5c5\ub370\uc774\ud2b8'),
]
for row, values in zip(doc.tables[5].rows, table5_data):
    for cell, value in zip(row.cells, values):
        cell.text = value

doc.save('temp.docx')
